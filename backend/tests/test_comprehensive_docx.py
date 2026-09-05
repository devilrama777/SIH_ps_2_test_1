import unittest
from pathlib import Path
from docx import Document
from backend.services.archive_service import HistoricalArchiveService
from backend.services.converter import MarkdownConverter
from backend.services.document_generator import DocumentGenerator
from backend import config


class TestComprehensiveDocxReporting(unittest.TestCase):
    """Test suite for comprehensive multi-source DOCX compilation and historical archives."""

    def setUp(self):
        self.archive_service = HistoricalArchiveService()
        self.converter = MarkdownConverter()
        self.doc_gen = DocumentGenerator()
        self.sample_csv = config.BASE_DIR / "default_data_backup" / "coal_production_report.csv"

    def test_historical_archive_service(self):
        """Verify that historical archives compute timelines, CAGR, and subsidiary revenue correctly."""
        summary = self.archive_service.generate_archive_intelligence_summary()
        self.assertIn("timeline", summary)
        self.assertIn("subsidiaries", summary)
        self.assertGreaterEqual(len(summary["timeline"]), 5, "Should have at least 5 years of production data")
        self.assertGreater(summary["cagr_5yr"], 0.0, "5-year CAGR should be positive")
        self.assertGreater(summary["total_subsidiary_revenue_crore"], 10000.0, "Total subsidiary revenue should exceed benchmark")

    def test_csv_and_image_conversion(self):
        """Verify multi-source conversion for tabular data and images."""
        if self.sample_csv.exists():
            csv_res = self.converter.convert(self.sample_csv)
            self.assertEqual(csv_res["file_type"], "csv")
            self.assertIn("National Coal Production", csv_res["markdown"])

        # Test bundle compilation
        bundle = self.converter.compile_multi_source_bundle([self.sample_csv])
        self.assertGreaterEqual(bundle["total_sources"], 1)
        self.assertIn("Integrated Multi-Source Intelligence Dossier", bundle["markdown"])

    def test_comprehensive_docx_generation(self):
        """Verify generation of Word DOCX with images, hyperlinks, and historical tables."""
        out_docx = self.doc_gen.generate_docx_report(
            template_name="executive_brief",
            report_id="TEST-SUITE-DOCX-001",
            summary_text="Unit test verification of multi-source DOCX compilation for Ministry of Coal.",
            sources_summary=[
                {"filename": "coal_production_report.csv", "type": "csv", "metadata": {"row_count": 26}}
            ],
            include_historical_archive=True
        )

        self.assertTrue(out_docx.exists(), "DOCX file must exist")
        self.assertGreater(out_docx.stat().st_size, 5000, "DOCX file must be non-empty")

        # Validate internal structure via python-docx
        doc = Document(str(out_docx))
        self.assertGreater(len(doc.paragraphs), 10, "DOCX should have rich paragraph structure")
        self.assertGreaterEqual(len(doc.tables), 4, "DOCX should contain at least 4 structured tables")

        # Check for OpenXML hyperlinks in paragraphs
        found_hyperlink = False
        for p in doc.paragraphs:
            for child in p._p:
                if child.tag.endswith("hyperlink"):
                    found_hyperlink = True
                    break
            if found_hyperlink:
                break
        self.assertTrue(found_hyperlink, "DOCX must contain active OpenXML clickable hyperlinks")


if __name__ == "__main__":
    unittest.main()
